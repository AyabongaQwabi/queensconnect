"""
CV tools: get/save CV doc (Firestore cvs/{waNumber}), generate PDF/DOCX and upload to Firebase Storage.
Used by cv_agent to collect CV data and produce downloadable files.
"""
import io
import logging
from datetime import timedelta, timezone
from typing import Optional

from google.adk.tools import FunctionTool

try:
    from ..config import FIREBASE_PROJECT_ID, FIREBASE_STORAGE_BUCKET
except ImportError:
    from config import FIREBASE_PROJECT_ID, FIREBASE_STORAGE_BUCKET

logger = logging.getLogger("queens_connect.tools.cv")

# Allowed top-level keys when saving CV doc (merge only these)
CV_TOP_LEVEL_KEYS = frozenset({
    "particulars", "education", "higherEducation", "workExperience",
    "bio", "professionalSummary", "coreSkills", "contact", "references",
    "fileLink", "format", "updatedAt",
})

# Minimum required for generation: must have these sections (content can be minimal)
REQUIRED_FOR_GENERATION = ("particulars", "education", "contact")


def _get_db():
    """Use Firestore from firebase_tools (ensures Firebase app is initialized)."""
    from .firebase_tools import _get_db as _firestore_db
    return _firestore_db()


def _get_storage_bucket():
    """Return the default Storage bucket for CV uploads."""
    bucket_name = FIREBASE_STORAGE_BUCKET or f"{FIREBASE_PROJECT_ID}.firebasestorage.app"
    try:
        from google.cloud import storage
        client = storage.Client(project=FIREBASE_PROJECT_ID)
        return client.bucket(bucket_name)
    except Exception as e:
        logger.exception("_get_storage_bucket failed: %s", e)
        raise


def _normalize_wa(wa_number: Optional[str]) -> str:
    if wa_number is None:
        return ""
    return str(wa_number).strip()


def _serialize_doc(data: dict) -> dict:
    """Convert Firestore timestamps to ISO strings for JSON-friendly return."""
    from datetime import datetime
    out = dict(data)
    for key in ("updatedAt", "createdAt"):
        if key in out and out[key] is not None:
            v = out[key]
            if hasattr(v, "seconds"):  # Firestore Timestamp
                out[key] = datetime.fromtimestamp(v.seconds + (getattr(v, "nanoseconds", 0) or 0) / 1e9, tz=timezone.utc).isoformat()
            elif hasattr(v, "isoformat"):
                out[key] = v.isoformat()
    return out


# ---------- Public tool implementations ----------


def get_cv_doc(wa_number: str) -> dict:
    """
    Get the CV document for this user. Call with waNumber from session state.
    Returns the document (including fileLink if a CV has been generated) or empty dict if none exists.
    Use on entry to decide: if fileLink is present, give the user the link; else start or continue collection.
    """
    wa = _normalize_wa(wa_number)
    if not wa:
        return {"exists": False, "error": "wa_number required"}
    try:
        db = _get_db()
        ref = db.collection("cvs").document(wa)
        doc = ref.get()
        if not doc.exists:
            return {"exists": False}
        data = doc.to_dict()
        out = _serialize_doc({"exists": True, "id": doc.id, **(data or {})})
        return out
    except Exception as e:
        logger.exception("get_cv_doc failed: %s", e)
        return {"exists": False, "error": str(e)}


def save_cv_doc(wa_number: str, data: dict) -> dict:
    """
    Save or merge CV data for this user. Call with waNumber from session and a dict containing
    any of: particulars, education, higherEducation, workExperience, bio, professionalSummary,
    coreSkills, contact, references. Optional particulars: driverLicenceCode, noticePeriod,
    expectedSalaryRange, workPermitStatus. higherEducation items may include fieldOfStudy.
    Used after each section or at end of collection. Merges into existing doc.
    Returns status and message.
    """
    wa = _normalize_wa(wa_number)
    if not wa:
        return {"status": "error", "error_message": "wa_number required"}
    # Only merge allowed top-level keys; do not allow fileLink/format to be set by agent (only generate_* set those)
    allowed = {k: v for k, v in (data or {}).items() if k in CV_TOP_LEVEL_KEYS and k not in ("fileLink", "format")}
    if not allowed:
        return {"status": "error", "error_message": "No valid CV fields to save. Use: particulars, education, higherEducation, workExperience, bio, professionalSummary, coreSkills, contact, references."}
    try:
        from google.cloud.firestore_v1 import SERVER_TIMESTAMP
    except ImportError:
        from google.cloud.firestore import SERVER_TIMESTAMP
    payload = {**allowed, "updatedAt": SERVER_TIMESTAMP}
    try:
        db = _get_db()
        ref = db.collection("cvs").document(wa)
        ref.set(payload, merge=True)
        logger.info("save_cv_doc wa=%s keys=%s", wa[:6] + "***", list(payload.keys()))
        return {"status": "success", "message": "CV data saved."}
    except Exception as e:
        logger.exception("save_cv_doc failed: %s", e)
        return {"status": "error", "error_message": str(e)}


def _validate_cv_for_generation(doc: dict) -> Optional[str]:
    """Return None if doc has all required sections; else return error message."""
    for key in REQUIRED_FOR_GENERATION:
        if key not in doc or doc[key] is None:
            return f"Missing required section: {key}. Please collect all sections before generating."
    p = doc.get("particulars") or {}
    if not isinstance(p, dict):
        return "particulars must be an object."
    if not (p.get("name") or p.get("surname")):
        return "Particulars must include at least name and surname."
    c = doc.get("contact") or {}
    if not isinstance(c, dict):
        return "contact must be an object."
    if not (c.get("email") or c.get("contactNumber")):
        return "Contact must include at least email or contact number."
    return None


def _build_pdf_bytes(doc: dict) -> bytes:
    """Render CV document to PDF bytes. Modern single-column template with navy accents and clean typography."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.lib.colors import HexColor
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.enums import TA_LEFT
    except ImportError:
        raise RuntimeError("reportlab is required for PDF generation. Install with: pip install reportlab")
    buffer = io.BytesIO()
    story = []
    # Design tokens: minimalist modern — navy headings, charcoal body (ATS-safe, print-friendly)
    NAVY = HexColor("#0A2342")
    CHARCOAL = HexColor("#333333")
    content_width = 160 * mm  # A4 210mm - 25mm margins each side

    styles = getSampleStyleSheet()
    # Name: 22pt bold, navy, prominent — enough space below so contact does not overlap
    name_style = ParagraphStyle(
        name="CVName",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=22,
        textColor=NAVY,
        spaceAfter=12,
        alignment=TA_LEFT,
    )
    # Contact line under name: 10pt charcoal, compact
    contact_style = ParagraphStyle(
        name="CVContact",
        parent=styles["Normal"],
        fontSize=10,
        textColor=CHARCOAL,
        spaceAfter=2,
        alignment=TA_LEFT,
    )
    # Section heading: 14pt bold navy, extra space before
    heading_style = ParagraphStyle(
        name="CVHeading",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=14,
        textColor=NAVY,
        spaceBefore=14,
        spaceAfter=6,
        alignment=TA_LEFT,
    )
    # Job title in work experience: 12pt bold charcoal
    job_title_style = ParagraphStyle(
        name="CVJobTitle",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=12,
        textColor=CHARCOAL,
        spaceAfter=2,
        alignment=TA_LEFT,
    )
    # Body / bullets: 11pt, 1.2 line spacing, charcoal
    body_style = ParagraphStyle(
        name="CVBody",
        parent=styles["Normal"],
        fontSize=11,
        leading=13,
        textColor=CHARCOAL,
        spaceAfter=4,
        alignment=TA_LEFT,
        leftIndent=0,
    )
    # Skills section: same as body with left indent for padding under "Core Skills" heading
    skills_body_style = ParagraphStyle(
        name="CVSkillsBody",
        parent=body_style,
        leftIndent=12,
    )
    # Subtext (company, date): 10pt
    sub_style = ParagraphStyle(
        name="CVSub",
        parent=styles["Normal"],
        fontSize=10,
        textColor=CHARCOAL,
        spaceAfter=4,
        alignment=TA_LEFT,
    )

    def escape(s: str) -> str:
        return (s or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    def add_para(text, style=None):
        if text is None:
            return
        if isinstance(text, list):
            text = "\n".join(str(x).strip() for x in text if str(x).strip())
        text = str(text or "").strip()
        if not text:
            return
        safe = escape(text)
        story.append(Paragraph(safe, style or body_style))
        story.append(Spacer(1, 4))

    p = doc.get("particulars") or {}
    c = doc.get("contact") or {}
    name = f"{p.get('name', '')} {p.get('surname', '')}".strip() or "CV"
    story.append(Paragraph(escape(name), name_style))
    contact_parts = []
    if c.get("email"):
        contact_parts.append(str(c.get("email", "")))
    if c.get("contactNumber"):
        contact_parts.append(str(c.get("contactNumber", "")))
    if c.get("address"):
        contact_parts.append(str(c.get("address", "")))
    if contact_parts:
        story.append(Spacer(1, 3))
        story.append(Paragraph(escape("  •  ".join(contact_parts)), contact_style))
    story.append(Paragraph(
        escape(f"DOB: {p.get('dateOfBirth', '')}  •  Nationality: {p.get('nationality', '')}  •  ID: {p.get('idNumber', '')}  •  Gender: {p.get('gender', '')}"),
        contact_style,
    ))
    if p.get("taxNumber"):
        story.append(Paragraph(escape(f"Tax number: {p.get('taxNumber')}"), contact_style))
    story.append(Paragraph(escape(f"Criminal record: {'Yes' if p.get('hasCriminalRecord') else 'No'}"), contact_style))
    # Horizontal rule: thin navy bar
    hr_table = Table([[""]], colWidths=[content_width], rowHeights=[2])
    hr_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), NAVY),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    story.append(hr_table)
    story.append(Spacer(1, 12))

    # Professional Summary
    summary = doc.get("professionalSummary") or doc.get("bio") or ""
    if summary:
        story.append(Paragraph("Professional Summary", heading_style))
        story.append(Spacer(1, 2))
        add_para(summary, body_style)
        story.append(Spacer(1, 6))

    # Core Skills — two-column pill-style list
    skills = doc.get("coreSkills") or []
    if skills and isinstance(skills, list):
        skill_strs = [str(s).strip() for s in skills if str(s).strip()]
        if skill_strs:
            story.append(Paragraph("Core Skills", heading_style))
            story.append(Spacer(1, 2))
            # Build two columns: left and right
            mid = (len(skill_strs) + 1) // 2
            left_col = "  •  ".join(skill_strs[:mid])
            right_col = "  •  ".join(skill_strs[mid:]) if mid < len(skill_strs) else ""
            if right_col:
                tbl = Table([[Paragraph(escape(left_col), body_style), Paragraph(escape(right_col), body_style)]], colWidths=[content_width / 2] * 2)
                tbl.setStyle(TableStyle([
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (0, -1), 12),
                    ("LEFTPADDING", (0, 1), (0, -1), 12),
                    ("RIGHTPADDING", (0, 0), (0, -1), 8),
                    ("RIGHTPADDING", (1, 0), (1, -1), 8),
                ]))
                story.append(tbl)
            else:
                add_para(left_col, skills_body_style)
            story.append(Spacer(1, 6))

    # Work Experience
    work = doc.get("workExperience") or []
    if work:
        story.append(Paragraph("Work Experience", heading_style))
        story.append(Spacer(1, 2))
        for w in work:
            if isinstance(w, dict):
                pos = w.get("position") or ""
                company = w.get("companyName") or ""
                year = w.get("year") or ""
                story.append(Paragraph(escape(pos), job_title_style))
                story.append(Paragraph(escape(f"{company} — {year}"), sub_style))
                desc = w.get("description", "")
                if isinstance(desc, list):
                    desc = "\n".join(str(x).strip() for x in desc if str(x).strip())
                if str(desc).strip():
                    add_para(desc, body_style)
                story.append(Spacer(1, 8))

    # Education
    story.append(Paragraph("Education", heading_style))
    story.append(Spacer(1, 2))
    e = doc.get("education") or {}
    add_para(f"High school: {e.get('highSchoolName', '')}. Highest grade: {e.get('highestGradePassed', '')}. Matriculated: {'Yes' + (' (' + str(e.get('matriculationYear', '')) + ')' if e.get('matriculationYear') else '') if e.get('matriculated') else 'No'}.", body_style)
    story.append(Spacer(1, 6))

    # Higher Education
    higher = doc.get("higherEducation") or []
    if higher:
        story.append(Paragraph("Higher Education", heading_style))
        story.append(Spacer(1, 2))
        for h in higher:
            if isinstance(h, dict):
                qual = h.get("degreeOrDiplomaOrCertificate", "") or ""
                inst = h.get("institution", "") or ""
                yr = h.get("yearPassed", "") or ""
                field = h.get("fieldOfStudy", "") or ""
                parts = [q for q in [qual, field] if q]
                line = " – ".join(parts) if parts else qual or "Qualification"
                if inst:
                    line += f" – {inst}"
                if yr:
                    line += f" ({yr})"
                add_para(line, body_style)
        story.append(Spacer(1, 6))

    # Additional Information
    extra = []
    if p.get("driverLicenceCode"):
        extra.append(f"Driver's licence: {p.get('driverLicenceCode')}")
    if p.get("noticePeriod"):
        extra.append(f"Notice period: {p.get('noticePeriod')}")
    if p.get("expectedSalaryRange"):
        extra.append(f"Expected salary: {p.get('expectedSalaryRange')}")
    if p.get("workPermitStatus"):
        extra.append(f"Work permit: {p.get('workPermitStatus')}")
    if extra:
        story.append(Paragraph("Additional Information", heading_style))
        story.append(Spacer(1, 2))
        add_para("  •  ".join(extra), body_style)
        story.append(Spacer(1, 6))

    # References
    refs = doc.get("references") or []
    if refs:
        story.append(Paragraph("References", heading_style))
        story.append(Spacer(1, 2))
        for r in refs:
            if isinstance(r, dict):
                add_para(f"{r.get('name', '')} – {r.get('relationship', '')}: {r.get('contactDetail', '')}", body_style)

    margin = 25 * mm
    pdf = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=margin,
        rightMargin=margin,
        topMargin=20 * mm,
        bottomMargin=20 * mm,
    )
    pdf.build(story)
    return buffer.getvalue()


def _build_docx_bytes(doc: dict) -> bytes:
    """Render CV document to DOCX bytes. Modern single-column template with navy accents and clean typography."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise RuntimeError("python-docx is required for DOCX generation. Install with: pip install python-docx")
    buffer = io.BytesIO()
    document = Document()
    # Margins: ~1 inch
    for section in document.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Navy #0A2342 and charcoal #333333 (RGB)
    NAVY_RGB = (10, 35, 66)
    CHARCOAL_RGB = (51, 51, 51)

    def add_heading_navy(text: str, level: int = 1):
        if not text:
            return
        p = document.add_heading(text, level=level)
        for run in p.runs:
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(*NAVY_RGB)

    def add_hr_navy():
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(12)
        p_border = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "0A2342")
        p_border.append(bottom)
        p._p.get_or_add_pPr().append(p_border)

    def add_para(text, font_size_pt=11, bold=False, color_rgb=None):
        if text is None:
            return
        if isinstance(text, list):
            text = "\n".join(str(x).strip() for x in text if str(x).strip())
        text = str(text or "").strip()
        if not text:
            return
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.size = Pt(font_size_pt)
        run.font.bold = bold
        if color_rgb:
            run.font.color.rgb = RGBColor(*color_rgb)
        else:
            run.font.color.rgb = RGBColor(*CHARCOAL_RGB)

    p = doc.get("particulars") or {}
    c = doc.get("contact") or {}
    name = f"{p.get('name', '')} {p.get('surname', '')}".strip() or "CV"
    name_para = document.add_paragraph()
    name_run = name_para.add_run(name)
    name_run.font.size = Pt(22)
    name_run.font.bold = True
    name_run.font.color.rgb = RGBColor(*NAVY_RGB)
    name_para.paragraph_format.space_after = Pt(4)

    contact_parts = []
    if c.get("email"):
        contact_parts.append(str(c.get("email", "")))
    if c.get("contactNumber"):
        contact_parts.append(str(c.get("contactNumber", "")))
    if c.get("address"):
        contact_parts.append(str(c.get("address", "")))
    if contact_parts:
        add_para("  •  ".join(contact_parts), font_size_pt=10)
    add_para(f"DOB: {p.get('dateOfBirth', '')}  •  Nationality: {p.get('nationality', '')}  •  ID: {p.get('idNumber', '')}  •  Gender: {p.get('gender', '')}", font_size_pt=10)
    if p.get("taxNumber"):
        add_para(f"Tax number: {p.get('taxNumber')}", font_size_pt=10)
    add_para(f"Criminal record: {'Yes' if p.get('hasCriminalRecord') else 'No'}", font_size_pt=10)
    add_hr_navy()

    summary = doc.get("professionalSummary") or doc.get("bio") or ""
    if summary:
        add_heading_navy("Professional Summary", level=1)
        document.paragraphs[-1].paragraph_format.space_before = Pt(14)
        add_para(summary, font_size_pt=11)

    skills = doc.get("coreSkills") or []
    if skills and isinstance(skills, list):
        skill_strs = [str(s).strip() for s in skills if str(s).strip()]
        if skill_strs:
            add_heading_navy("Core Skills", level=1)
            document.paragraphs[-1].paragraph_format.space_before = Pt(14)
            mid = (len(skill_strs) + 1) // 2
            left_col = "  •  ".join(skill_strs[:mid])
            right_col = "  •  ".join(skill_strs[mid:]) if mid < len(skill_strs) else ""
            tbl = document.add_table(rows=1, cols=2)
            tbl.cell(0, 0).text = left_col
            tbl.cell(0, 1).text = right_col
            for cell in tbl.columns[0].cells:
                cell.paragraphs[0].paragraph_format.space_after = Pt(2)
            for cell in tbl.columns[1].cells:
                cell.paragraphs[0].paragraph_format.space_after = Pt(2)
            for row in tbl.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for r in para.runs:
                            r.font.size = Pt(11)
                            r.font.color.rgb = RGBColor(*CHARCOAL_RGB)
            document.add_paragraph()

    work = doc.get("workExperience") or []
    if work:
        add_heading_navy("Work Experience", level=1)
        document.paragraphs[-1].paragraph_format.space_before = Pt(14)
        for w in work:
            if isinstance(w, dict):
                document.add_paragraph()
                add_para(w.get("position", ""), font_size_pt=12, bold=True)
                add_para(f"{w.get('companyName', '')} — {w.get('year', '')}", font_size_pt=10)
                desc = w.get("description", "")
                if isinstance(desc, list):
                    desc = "\n".join(str(x).strip() for x in desc if str(x).strip())
                if str(desc).strip():
                    add_para(desc, font_size_pt=11)
                document.add_paragraph()

    add_heading_navy("Education", level=1)
    document.paragraphs[-1].paragraph_format.space_before = Pt(14)
    e = doc.get("education") or {}
    add_para(f"High school: {e.get('highSchoolName', '')}. Highest grade: {e.get('highestGradePassed', '')}. Matriculated: {'Yes' + (' (' + str(e.get('matriculationYear', '')) + ')' if e.get('matriculationYear') else '') if e.get('matriculated') else 'No'}.", font_size_pt=11)

    higher = doc.get("higherEducation") or []
    if higher:
        add_heading_navy("Higher Education", level=1)
        document.paragraphs[-1].paragraph_format.space_before = Pt(14)
        for h in higher:
            if isinstance(h, dict):
                qual = h.get("degreeOrDiplomaOrCertificate", "") or ""
                inst = h.get("institution", "") or ""
                yr = h.get("yearPassed", "") or ""
                field = h.get("fieldOfStudy", "") or ""
                parts = [q for q in [qual, field] if q]
                line = " – ".join(parts) if parts else qual or "Qualification"
                if inst:
                    line += f" – {inst}"
                if yr:
                    line += f" ({yr})"
                add_para(line, font_size_pt=11)

    extra = []
    if p.get("driverLicenceCode"):
        extra.append(f"Driver's licence: {p.get('driverLicenceCode')}")
    if p.get("noticePeriod"):
        extra.append(f"Notice period: {p.get('noticePeriod')}")
    if p.get("expectedSalaryRange"):
        extra.append(f"Expected salary: {p.get('expectedSalaryRange')}")
    if p.get("workPermitStatus"):
        extra.append(f"Work permit: {p.get('workPermitStatus')}")
    if extra:
        add_heading_navy("Additional Information", level=1)
        document.paragraphs[-1].paragraph_format.space_before = Pt(14)
        add_para("  •  ".join(extra), font_size_pt=11)

    refs = doc.get("references") or []
    if refs:
        add_heading_navy("References", level=1)
        document.paragraphs[-1].paragraph_format.space_before = Pt(14)
        for r in refs:
            if isinstance(r, dict):
                add_para(f"{r.get('name', '')} – {r.get('relationship', '')}: {r.get('contactDetail', '')}", font_size_pt=11)

    document.save(buffer)
    return buffer.getvalue()


def _upload_cv_and_get_url(wa_number: str, file_bytes: bytes, extension: str) -> str:
    """Upload bytes to cvs/{wa_number}/cv.{extension}, return signed URL (48h)."""
    wa = _normalize_wa(wa_number)
    bucket = _get_storage_bucket()
    blob_path = f"cvs/{wa}/cv.{extension}"
    blob = bucket.blob(blob_path)
    blob.upload_from_string(file_bytes, content_type="application/pdf" if extension == "pdf" else "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    try:
        url = blob.generate_signed_url(
            version="v4",
            expiration=timedelta(hours=48),
            method="GET",
        )
        return url
    except Exception as e:
        logger.warning("generate_signed_url failed (%s), using public URL if available", e)
        return blob.public_url or blob_path


def generate_cv_pdf(wa_number: str) -> dict:
    """
    Generate a PDF CV from the user's saved CV data, upload it to Firebase Storage,
    update the CV doc with the download link, and return the link.
    Call only after all sections have been collected and save_cv_doc has been used.
    Returns { status, message, fileLink } or { status: "error", error_message: "..." }.
    """
    wa = _normalize_wa(wa_number)
    if not wa:
        return {"status": "error", "error_message": "wa_number required"}
    try:
        db = _get_db()
        ref = db.collection("cvs").document(wa)
        doc_snap = ref.get()
        if not doc_snap or not doc_snap.exists:
            return {"status": "error", "error_message": "No CV data found. Please complete the CV collection first."}
        doc = doc_snap.to_dict() or {}
        err = _validate_cv_for_generation(doc)
        if err:
            return {"status": "error", "error_message": err}
        pdf_bytes = _build_pdf_bytes(doc)
        file_link = _upload_cv_and_get_url(wa, pdf_bytes, "pdf")
        try:
            from google.cloud.firestore_v1 import SERVER_TIMESTAMP
        except ImportError:
            from google.cloud.firestore import SERVER_TIMESTAMP
        ref.set({"fileLink": file_link, "format": "pdf", "updatedAt": SERVER_TIMESTAMP}, merge=True)
        logger.info("generate_cv_pdf wa=%s success", wa[:6] + "***")
        return {"status": "success", "message": "Your CV (PDF) is ready.", "fileLink": file_link}
    except Exception as e:
        logger.exception("generate_cv_pdf failed: %s", e)
        return {"status": "error", "error_message": str(e)}


def generate_cv_docx(wa_number: str) -> dict:
    """
    Generate a Microsoft Word (DOCX) CV from the user's saved CV data, upload it to Firebase Storage,
    update the CV doc with the download link, and return the link.
    Call only after all sections have been collected. Returns { status, message, fileLink } or error.
    """
    wa = _normalize_wa(wa_number)
    if not wa:
        return {"status": "error", "error_message": "wa_number required"}
    try:
        db = _get_db()
        ref = db.collection("cvs").document(wa)
        doc_snap = ref.get()
        if not doc_snap or not doc_snap.exists:
            return {"status": "error", "error_message": "No CV data found. Please complete the CV collection first."}
        doc = doc_snap.to_dict() or {}
        err = _validate_cv_for_generation(doc)
        if err:
            return {"status": "error", "error_message": err}
        docx_bytes = _build_docx_bytes(doc)
        file_link = _upload_cv_and_get_url(wa, docx_bytes, "docx")
        try:
            from google.cloud.firestore_v1 import SERVER_TIMESTAMP
        except ImportError:
            from google.cloud.firestore import SERVER_TIMESTAMP
        ref.set({"fileLink": file_link, "format": "docx", "updatedAt": SERVER_TIMESTAMP}, merge=True)
        logger.info("generate_cv_docx wa=%s success", wa[:6] + "***")
        return {"status": "success", "message": "Your CV (Word) is ready.", "fileLink": file_link}
    except Exception as e:
        logger.exception("generate_cv_docx failed: %s", e)
        return {"status": "error", "error_message": str(e)}


# ---------- ADK FunctionTool wrappers ----------
get_cv_doc_tool = FunctionTool(get_cv_doc)
save_cv_doc_tool = FunctionTool(save_cv_doc)
generate_cv_pdf_tool = FunctionTool(generate_cv_pdf)
generate_cv_docx_tool = FunctionTool(generate_cv_docx)
